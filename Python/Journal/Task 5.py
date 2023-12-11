# Importing modules and read databases into a dataframe
import docx
import sqlite3
import pandas as pd
from docx.shared import Inches
import json
from datetime import datetime, date

con1 = sqlite3.connect('DATA.sqlite')
df = pd.read_sql_query("SELECT * from Data", con1)
con2 = sqlite3.connect('PREDICTION.sqlite')
df2 = pd.read_sql_query("SELECT * from Predictions", con2)
df['Prediction'] = df2['Prediction']

# Docx operations

last_row = df.iloc[-1]
date1=last_row.Date
dtime = datetime.strptime(date1,'%Y-%m-%d')
heading = dtime.strftime('%A, %d. %B %Y')

# Current date
today = date.today()
date = today.strftime('%Y-%m-%d')

# Try to open a journal docx. if it is not available, create a new one
isNew = False
try:
    doc = docx.Document('Journal.docx')
except:
    doc = docx.Document()
    doc.save('Journal.docx')
    isnew = True
doc = docx.Document('Journal.docx')


# Function to add first paragraph
def addParagraph(doc, last_row, summary='...'):
    sleep_time = f'Slept: {last_row.Went_to_bed} - {last_row.Woke_up}'
    hours, minutes, seconds = list(map(int, last_row.Sleeping_hours.split(':')))
    slept_time = f'Sleep time: {hours}H {minutes}m'
    cigarretes = f'Cigarettes: {last_row.Cigarettes}'
    sports = f'Sports: {last_row.Sport}'

    weather = f'Weather: {last_row.Weather}'
    myday = f'My day: {last_row.DayAnswer} day'
    prognosis = f'Prognosis: {last_row.Prediction} day'
    new_line = '\n'

    para = f'{sleep_time}{new_line}{slept_time}{new_line}{cigarretes}{new_line}{sports}{new_line}{weather}{new_line}{myday}{new_line}{prognosis}{new_line}'

    doc.add_paragraph(summary)

    doc.add_paragraph(para)


# Function to add graphs

def addGraphs(doc, date):
    alcohol_img = f'{date}_Alcohol.jpg'
    cig_img = f'{date}_cigarettes.jpg'
    tbl = doc.add_table(rows=1, cols=2)
    row_cells = tbl.add_row().cells
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(alcohol_img, width=Inches(3.1))
    paragraph = row_cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(cig_img, width=Inches(3.1))


# Function to add photos
def addPhotos(doc, photos):
    tbl = doc.add_table(rows=1, cols=2)
    row_cells = tbl.add_row().cells
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(photos[0], width=Inches(2))
    paragraph = row_cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(photos[1], width=Inches(2))


# Combining
canWrite = False
if isNew or len(doc.paragraphs) == 0:
    canWrite = True
elif doc.paragraphs[-4].text != heading:
    canWrite = True
if canWrite:

    doc.add_heading(heading, 0)

    addParagraph(doc, last_row=last_row)
    photos = [f'{date}_1.jpg', f'{date}_2.jpg']
    addPhotos(doc, photos)
    addGraphs(doc, date)

    doc.add_page_break()
    print('Document updated')
else:
    print('Document is already upto date')
doc.save('Journal.docx')

